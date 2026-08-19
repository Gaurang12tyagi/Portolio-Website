# -*- coding: utf-8 -*-
"""Generate the ATS resume .docx and .txt from a single content source.
The .pdf is produced separately from resume.html via headless Chrome."""
import io, os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "/Users/gaurang/Downloads"

NAME = "GAURANG TYAGI"
TITLE = "Frontend Developer (React.js) | Next.js, TypeScript, Redux | Full Stack MERN"
CONTACT = [
    "Jaipur, Rajasthan, India 302012 | +91 97856 06965 | gaurangtyagi95@gmail.com",
    "LinkedIn: linkedin.com/in/gaurang-tyagi12 | GitHub: github.com/Gaurang12tyagi",
    "Portfolio: portolio-website-p5tr-gaurang-tyagis-projects.vercel.app",
]

SUMMARY = (
    "Frontend-focused Full Stack Developer with 3+ years of experience designing and shipping "
    "production-grade user interfaces with React.js, Next.js, and TypeScript. Specialized in "
    "component-driven architecture, reusable design systems, state management with Redux Toolkit "
    "and Context API, server-side rendering, and frontend performance engineering through code "
    "splitting, lazy loading, virtualization, and memoization. Backed by solid Node.js, Express.js, "
    "and microservices experience for end-to-end feature delivery. Proven record of improving page "
    "load speed by 15%, raising mobile UX scores by 25%, and cutting API response times by 60% "
    "through Redis caching. Strong collaborator in Agile/Scrum teams with experience in code "
    "reviews, accessibility standards, and mentoring junior developers."
)

SKILLS = [
    ("Frontend Development",
     "React.js, Next.js, TypeScript, Redux, Redux Toolkit, Context API, React Hooks and Custom Hooks, "
     "React Router, React Query, React Hook Form, Server-Side Rendering (SSR), Static Site Generation (SSG), "
     "Incremental Static Regeneration (ISR), Component-Based Architecture, Reusable Component Libraries and "
     "Design Systems, Tailwind CSS, Material UI, SCSS, Framer Motion, Responsive and Mobile-First Design, "
     "Cross-Browser Compatibility, Web Accessibility (WCAG), Code Splitting, Lazy Loading, List Virtualization, "
     "Memoization, Error Boundaries, DOM Manipulation, AJAX"),
    ("Programming Languages",
     "JavaScript (ES6+), TypeScript, C++, SQL, HTML5, CSS3"),
    ("Backend Development",
     "Node.js, Express.js, RESTful API Design and Development, WebSocket API, Socket.io, Microservices "
     "Architecture, Server-Side Logic, Cron Jobs and Job Scheduling, Role-Based Access Control (RBAC), "
     "Authentication and Authorization, Third-Party API Integration, Error Handling, Stripe Payment Gateway "
     "Integration"),
    ("Databases",
     "MongoDB, PostgreSQL, MySQL, Elasticsearch, Redis, Database Design and Data Modeling, Query Optimization, "
     "Database Indexing, Stored Procedures, Complex Joins, Aggregation Pipelines, Schema Design"),
    ("DevOps and Tooling",
     "Git, GitHub, GitHub Actions, CI/CD (Continuous Integration and Continuous Deployment), Docker, "
     "Docker Compose, Multi-Stage Builds, Containerization, Webpack, Babel, ESLint, Prettier, Postman, "
     "Chrome DevTools, Lighthouse, npm, Vercel, VS Code, JIRA"),
    ("Engineering Practices",
     "Agile, Scrum, Sprint Planning, Code Review, Unit Testing, Debugging and Troubleshooting, Frontend "
     "Performance Optimization, Caching Strategies, State Management, API Documentation, Version Control, "
     "Build Pipelines, Code Quality Standards, Software Development Life Cycle (SDLC), Object-Oriented "
     "Programming (OOP), Data Structures and Algorithms"),
    ("Soft Skills",
     "Problem Solving, Cross-Functional Collaboration, Technical Communication, Mentoring, Ownership, "
     "Time Management"),
]

EXPERIENCE = [
    {
        "title": u"Software Engineer — Sarv.com",
        "meta": u"Jaipur, Rajasthan, India | May 2025 – Present",
        "bullets": [
            "Built and shipped React.js and Next.js frontends for full stack web applications serving thousands of daily active users, owning UI architecture, routing, and state management end to end.",
            "Engineered scalable RESTful APIs with Node.js and Express.js handling high-volume daily traffic, sustaining optimized response times under 200 ms.",
            "Implemented advanced caching strategies and database indexing across PostgreSQL and MongoDB, reducing query execution time by 40%.",
            "Designed and integrated WebSocket APIs and their React real-time UI layer for live communication features, increasing user engagement by 30%.",
            "Built GitHub Actions CI/CD pipelines (ESLint, automated tests, Webpack build, deploy), cutting release cycles by 35%.",
            "Containerized services with Docker using multi-stage builds and Docker Compose, ensuring environment parity across development and production.",
            "Collaborated with cross-functional teams across Agile sprints to deliver features on schedule, maintaining 95% sprint velocity.",
        ],
    },
    {
        "title": u"Software Developer — Grape Technology",
        "meta": u"Dubai, United Arab Emirates | May 2023 – May 2025",
        "bullets": [
            "Developed and maintained responsive, mobile-first user interfaces with React.js, improving mobile user experience scores by 25%.",
            "Built 20+ reusable React components in TypeScript, reducing code redundancy by 20% and accelerating development cycles.",
            "Implemented Redux and Context API for state management across complex applications, reducing state-related defects by 25%.",
            "Configured Webpack 5 (code splitting, lazy loading, tree shaking) with Babel transpilation and React memoization, achieving 15% faster page load times.",
            "Enforced code quality with custom ESLint rule sets integrated into GitHub Actions checks on every pull request.",
            "Integrated role-based authentication and protected routing for 3+ user roles, strengthening application security and access control.",
            "Partnered with backend developers to design and consume RESTful APIs, ensuring reliable data flow and robust error handling.",
        ],
    },
]

PROJECTS = [
    {
        "title": u"Octopus Social — Social Commerce, Events and Service Booking Platform",
        "stack": "React.js, Next.js, TypeScript, Redux Toolkit, React Query, Tailwind CSS, Socket.io, Node.js, MongoDB, Stripe API",
        "bullets": [
            "Led frontend architecture and development of Octopus Social, a unified platform that integrates a social media network with e-commerce, online and physical event management, and service booking, delivered as a modular React.js and Next.js application.",
            "Designed a shared component library of 60+ reusable TypeScript React components and a token-based design system, cutting new-feature UI development time by 35% and removing duplicate styling across five product modules.",
            "Built an infinite-scroll social feed with list virtualization, optimistic UI updates, and React Query caching, keeping interactions smooth at 10,000+ posts and reducing unnecessary re-renders by 45%.",
            "Developed shoppable-post and storefront journeys covering product listing, faceted filters, cart, and Stripe checkout, using Redux Toolkit for global state with loading and failure states handled through React Error Boundaries and Suspense.",
            "Implemented the event management module for both online and physical events, including calendar views, ticket booking, seat and slot selection, attendee dashboards, and QR-based check-in screens.",
            "Created a multi-step service booking flow with an availability calendar, dynamic slot picker, and real-time Socket.io notifications for booking confirmation and status updates.",
            "Optimized delivery with Next.js SSR and ISR, route-level code splitting, lazy loading, and memoization, achieving a Lighthouse performance score above 95 and Largest Contentful Paint under 2 seconds on fully responsive, WCAG-compliant screens.",
        ],
    },
    {
        "title": "E-Commerce Platform with Real-Time Features",
        "stack": "React.js, Node.js, Express.js, MongoDB, Socket.io, Elasticsearch, Stripe API",
        "bullets": [
            "Built a full-featured e-commerce platform with a React.js storefront and real-time inventory updates over WebSocket, supporting high concurrent user loads.",
            "Integrated the Stripe payment gateway with a secure checkout flow and comprehensive error handling for reliable transactions.",
            "Implemented advanced filtering, sorting, and search UI backed by Elasticsearch, improving search accuracy by 40%.",
        ],
    },
    {
        "title": "Task Management System with Microservices",
        "stack": "Next.js, React.js, Redux, Express.js, PostgreSQL, Redis, Docker",
        "bullets": [
            "Architected a microservices-based task management system with independent services for authentication, notifications, and tasks, fronted by a Next.js dashboard with drag-and-drop boards and role-based views.",
            "Implemented a Redis caching layer that reduced database load by 50% and improved API response times by 60%.",
            "Deployed containerized services with Docker, enabling straightforward scaling and simplified, repeatable deployments.",
        ],
    },
]

EDUCATION = [
    ("Master of Computer Applications (MCA), Computer Science",
     u"Lovely Professional University | September 2022 – September 2024 | 88%"),
    ("Bachelor of Computer Applications (BCA), Computer Science",
     u"University of Rajasthan | August 2019 – August 2022 | 80%"),
]

CERTIFICATIONS = [
    u"JavaScript Algorithms and Data Structures — freeCodeCamp",
    u"Responsive Web Design — freeCodeCamp",
    u"C++ Certification — Cipher School",
    u"TCS ESG Virtual Experience Program — Tata Consultancy Services",
]

ACHIEVEMENTS = [
    "Received the Employee of the Quarter award at Grape Technology for outstanding engineering performance and code quality.",
    "Mentored 3 junior developers on React and Node.js best practices, improving team productivity and code quality standards.",
]


# ---------------------------------------------------------------- plain text
def build_txt(path):
    L = []
    L.append(NAME)
    L.append(TITLE)
    L.extend(CONTACT)
    L.append("")
    L.append("PROFESSIONAL SUMMARY")
    L.append(SUMMARY)
    L.append("")
    L.append("TECHNICAL SKILLS")
    for label, body in SKILLS:
        L.append(u"%s: %s" % (label, body))
    L.append("")
    L.append("PROFESSIONAL EXPERIENCE")
    for job in EXPERIENCE:
        L.append(job["title"])
        L.append(job["meta"])
        for b in job["bullets"]:
            L.append(u"\t•\t" + b)
        L.append("")
    L.append("PROJECTS")
    for pr in PROJECTS:
        L.append(pr["title"])
        L.append(pr["stack"])
        for b in pr["bullets"]:
            L.append(u"\t•\t" + b)
        L.append("")
    L.append("EDUCATION")
    for deg, meta in EDUCATION:
        L.append(deg)
        L.append(meta)
        L.append("")
    L.append("CERTIFICATIONS")
    for c in CERTIFICATIONS:
        L.append(u"\t•\t" + c)
    L.append("")
    L.append("ACHIEVEMENTS")
    for a in ACHIEVEMENTS:
        L.append(u"\t•\t" + a)
    io.open(path, "w", encoding="utf-8").write(u"\n".join(L) + u"\n")


# --------------------------------------------------------------------- docx
BODY_PT = 9.5
HEAD_PT = 11.0


def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    pf = st.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.05


def _p(doc, text="", bold=False, italic=False, size=BODY_PT, after=2, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Arial"
    return p


def _bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    pPr.append(borders)


def _heading(doc, text):
    p = _p(doc, text.upper(), bold=True, size=HEAD_PT, after=4, before=8)
    _bottom_border(p)
    return p


def _label_para(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.size = Pt(BODY_PT)
    r.font.name = "Arial"
    r2 = p.add_run(body)
    r2.font.size = Pt(BODY_PT)
    r2.font.name = "Arial"
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    r = p.add_run(text)
    r.font.size = Pt(BODY_PT)
    r.font.name = "Arial"
    return p


def build_docx(path):
    doc = docx.Document()
    _style(doc)
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)  # A4
    s.left_margin = s.right_margin = Inches(0.5)
    s.top_margin = s.bottom_margin = Inches(0.5)

    _p(doc, NAME, bold=True, size=21, after=3)
    _p(doc, TITLE, bold=True, size=10, after=4)
    for line in CONTACT:
        _p(doc, line, after=1)

    _heading(doc, "Professional Summary")
    _p(doc, SUMMARY, after=3)

    _heading(doc, "Technical Skills")
    for label, body in SKILLS:
        _label_para(doc, label, body)

    _heading(doc, "Professional Experience")
    for job in EXPERIENCE:
        _p(doc, job["title"], bold=True, after=0, before=4)
        _p(doc, job["meta"], italic=True, after=2)
        for b in job["bullets"]:
            _bullet(doc, b)

    _heading(doc, "Projects")
    for pr in PROJECTS:
        _p(doc, pr["title"], bold=True, after=0, before=4)
        _p(doc, pr["stack"], italic=True, after=2)
        for b in pr["bullets"]:
            _bullet(doc, b)

    _heading(doc, "Education")
    for deg, meta in EDUCATION:
        _p(doc, deg, bold=True, after=0, before=3)
        _p(doc, meta, italic=True, after=2)

    _heading(doc, "Certifications")
    for c in CERTIFICATIONS:
        _bullet(doc, c)

    _heading(doc, "Achievements")
    for a in ACHIEVEMENTS:
        _bullet(doc, a)

    doc.save(path)


if __name__ == "__main__":
    build_txt(os.path.join(OUT_DIR, "Gaurang_Tyagi_Resume_ATS.txt"))
    build_docx(os.path.join(OUT_DIR, "Gaurang_Tyagi_Resume_ATS.docx"))
    print("txt + docx written to", OUT_DIR)
